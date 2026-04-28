from pathlib import Path

base = Path(r'c:\Users\lmthf\OneDrive - ISCTE-IUL\Desktop\Projetos\Inqu_Novos_Estudantes')
docx = base / 'Inquérito aos Novos Alunos do 3ºCiclo (25-26) - Versão Final_21-04-2025.docx'
xls = base / 'Cópia_bd_selo.XLS'
report = base / '_source_report.txt'

lines = []
lines.append(f'DOCX exists: {docx.exists()}')
lines.append(f'XLS exists: {xls.exists()}')

try:
    from docx import Document
    doc = Document(str(docx))
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    lines.append('')
    lines.append('DOCX paragraphs (first 150):')
    for i, p in enumerate(paras[:150], 1):
        lines.append(f'{i:03d}: {p}')
    lines.append('')
    lines.append('DOCX tables:')
    for ti, table in enumerate(doc.tables, 1):
        lines.append(f'Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols')
        for ri, row in enumerate(table.rows[:10], 1):
            cells = [c.text.strip().replace('\n', ' | ') for c in row.cells]
            lines.append(f'  Row {ri}: {cells}')
except Exception as e:
    lines.append(f'DOCX error: {type(e).__name__} {e}')

try:
    import pandas as pd
    xl = pd.ExcelFile(str(xls))
    lines.append('')
    lines.append(f'XLS sheets: {xl.sheet_names}')
    for sheet in xl.sheet_names:
        try:
            df = pd.read_excel(str(xls), sheet_name=sheet)
            lines.append('')
            lines.append(f'Sheet {sheet}: shape={df.shape}')
            lines.append(df.head(12).to_string())
        except Exception as e:
            lines.append(f'Could not read sheet {sheet}: {type(e).__name__} {e}')
except Exception as e:
    lines.append(f'XLS error: {type(e).__name__} {e}')

report.write_text('\n'.join(lines), encoding='utf-8')
print(report)
